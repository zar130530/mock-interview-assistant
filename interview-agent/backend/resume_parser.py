"""简历解析：支持 PDF / Word(.docx) / 纯文本，按内容哈希去重。"""
import hashlib
import io


def parse_resume(file_bytes: bytes, filename: str) -> str:
    ext = filename.lower().rsplit(".", 1)[-1] if "." in filename else ""
    if ext == "pdf":
        return _parse_pdf(file_bytes)
    if ext == "docx":
        return _parse_docx(file_bytes)
    # txt 及其它按文本处理
    return file_bytes.decode("utf-8", errors="ignore")


def content_hash(text: str) -> str:
    return hashlib.sha256(text.strip().encode("utf-8")).hexdigest()


def _parse_pdf(file_bytes: bytes) -> str:
    import pdfplumber
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            parts.append(page.extract_text() or "")
    return "\n".join(parts).strip()


def _parse_docx(file_bytes: bytes) -> str:
    from docx import Document
    doc = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts).strip()
